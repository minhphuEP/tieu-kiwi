"""Extract plain text from a source document for LLM ingestion.

Supported formats:
  .md / .txt / .markdown  — read as UTF-8
  .pdf                    — pypdf (page-joined text)
  .docx                   — python-docx (paragraph-joined text)
  .doc                    — macOS `textutil` fallback (Linux/Windows: not supported)

All extractors return a single string. Formatting is preserved best-effort but
LLMs don't care about layout — they need the words. Blank lines are collapsed
by the caller if desired.
"""
from __future__ import annotations

import os
import subprocess
import tempfile
from pathlib import Path


class UnsupportedFormatError(RuntimeError):
    pass


SUPPORTED_EXTS = {".md", ".markdown", ".txt", ".pdf", ".docx", ".doc"}


def read_text(path):
    """Auto-detect the file type and return its plain text.

    Raises UnsupportedFormatError if the extension is unknown.
    Raises RuntimeError with a clear message if the extractor for a known
    extension is not installed.
    """
    p = Path(path)
    ext = p.suffix.lower()
    if ext in (".md", ".markdown", ".txt"):
        return p.read_text(encoding="utf-8", errors="replace")
    if ext == ".pdf":
        return _read_pdf(p)
    if ext == ".docx":
        return _read_docx(p)
    if ext == ".doc":
        return _read_doc(p)
    raise UnsupportedFormatError(
        f"Unsupported extension {ext!r} for {path}. "
        f"Supported: {sorted(SUPPORTED_EXTS)}"
    )


# --- individual extractors -------------------------------------------------

def _read_pdf(path):
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise RuntimeError(
            "PDF support requires `pypdf`. Install with: pip install pypdf"
        ) from e
    reader = PdfReader(str(path))
    parts = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception as e:
            text = f"[warn: page {i+1} extract failed: {e}]"
        parts.append(text)
    return "\n\n".join(parts).strip()


def _read_docx(path):
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError(
            "DOCX support requires `python-docx`. Install with: pip install python-docx"
        ) from e
    doc = Document(str(path))
    # Grab paragraphs + tables so we don't drop tabular acceptance criteria.
    parts = []
    for para in doc.paragraphs:
        if para.text:
            parts.append(para.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def _read_doc(path):
    """Legacy .doc via macOS `textutil`. Not portable to Linux/Windows.

    On non-macOS, ask the user to save as .docx / .pdf first.
    """
    which = subprocess.run(["which", "textutil"], capture_output=True, text=True)
    if which.returncode != 0:
        raise RuntimeError(
            f"Cannot read {path}: `textutil` not found (macOS only). "
            f"Please convert to .docx or .pdf and retry."
        )
    with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as tf:
        out = tf.name
    try:
        subprocess.run(
            ["textutil", "-convert", "txt", str(path), "-output", out],
            check=True, capture_output=True,
        )
        return Path(out).read_text(encoding="utf-8", errors="replace").strip()
    finally:
        try: os.unlink(out)
        except OSError: pass


if __name__ == "__main__":
    import sys
    if len(sys.argv) != 2:
        print("Usage: python -m tieukiwi.text_extract <file>")
        sys.exit(1)
    print(read_text(sys.argv[1]))
